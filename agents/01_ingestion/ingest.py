"""
╔══════════════════════════════════════════════════════════════════╗
║  AGENT 1 — INGESTION                                            ║
║                                                                  ║
║  Multi-format manuscript ingestion.                             ║
║  Supported input formats: .docx, .doc, .pdf, .epub, .txt, .md  ║
║                                                                  ║
║  Also handles single-file manuscripts: if one big file          ║
║  contains the whole book, it auto-splits into chapters by        ║
║  detecting heading styles, "Chapter N" patterns, etc.           ║
║                                                                  ║
║  Input:   input/chapters/*  (any supported format)             ║
║  Output:  output/ingested/chapter_XX.json                       ║
║           output/ingested/images/  (embedded images, DOCX only) ║
║           output/ingested/ingestion_summary.json                ║
║                                                                  ║
║  How to run:                                                     ║
║    python agents/01_ingestion/ingest.py                         ║
║                                                                  ║
║  NOTE: Pure file-in / file-out. No database imports.            ║
║        The runner (runner.py) handles DB updates after this     ║
║        script completes.                                         ║
╚══════════════════════════════════════════════════════════════════╝
"""

import json
import os
import re
import zipfile
from pathlib import Path
from datetime import datetime

# python-docx: reads .docx/.doc files
from docx import Document

# colorama: coloured terminal output
from colorama import init, Fore, Style
init(autoreset=True)


# ── Optional dependencies ──────────────────────────────────────────────────────
# We import these lazily inside each reader so the agent still works even if
# only some packages are installed.  Missing deps produce a clear warning
# rather than a crash at startup.

def _import_pdfplumber():
    try:
        import pdfplumber
        return pdfplumber
    except ImportError:
        return None

def _import_epub_libs():
    """Returns (ebooklib_module, BeautifulSoup_class) or (None, None)."""
    try:
        import ebooklib
        from ebooklib import epub as epub_mod
        from bs4 import BeautifulSoup
        return ebooklib, epub_mod, BeautifulSoup
    except ImportError:
        return None, None, None


# ── Where things live ──────────────────────────────────────────────────────────
BASE_DIR = (
    Path(os.environ["CASSIAN_PROJECT_DIR"])
    if "CASSIAN_PROJECT_DIR" in os.environ
    else Path(__file__).resolve().parent.parent.parent
)
INPUT_DIR  = BASE_DIR / "input"  / "chapters"
OUTPUT_DIR = BASE_DIR / "output" / "ingested"
IMAGES_DIR = OUTPUT_DIR / "images"


# ── Terminal helpers ───────────────────────────────────────────────────────────
def ok(msg):   print(f"{Fore.GREEN}  ✓ {msg}{Style.RESET_ALL}")
def info(msg): print(f"{Fore.CYAN}  → {msg}{Style.RESET_ALL}")
def warn(msg): print(f"{Fore.YELLOW}  ⚠ {msg}{Style.RESET_ALL}")
def err(msg):  print(f"{Fore.RED}  ✗ {msg}{Style.RESET_ALL}")


# ── Chapter-break detection patterns ──────────────────────────────────────────
# Used to recognise chapter headings inside a single-file manuscript.
# Conservative — requires the marker to be at the START of the paragraph
# so we don't accidentally split on "Chapter and verse" mid-sentence.
_CHAPTER_RE = re.compile(
    r"^\s*"
    r"("
    r"chapter\s+(\d+|[ivxlcdmIVXLCDM]+)"   # Chapter 1 / Chapter IV
    r"|part\s+(\d+|[ivxlcdmIVXLCDM]+)"     # Part 1 / Part II
    r"|\d+\s*$"                              # bare number on its own line
    r"|\d+\."                               # 1. Title
    r")",
    re.IGNORECASE,
)

def _is_chapter_marker(text: str) -> bool:
    return bool(_CHAPTER_RE.match(text.strip()))


# ── Filename ↔ chapter ID helpers (unchanged from v1) ─────────────────────────

def _strip_order_prefix(stem: str) -> str:
    """Remove the leading '01_' style prefix the runner adds."""
    return re.sub(r"^\d+_", "", stem)


def extract_chapter_number(filename: str) -> int | None:
    stem = Path(filename).stem
    prefix = re.match(r"^(\d+)_", stem)
    if prefix:
        return int(prefix.group(1))
    match = re.search(r"(\d+)", filename)
    return int(match.group(1)) if match else None


def extract_chapter_id(filename: str) -> str:
    stem = Path(filename).stem.lower()
    if "epilogue" in stem:
        return "epilogue"
    prefix = re.match(r"^(\d+)_", stem)
    if prefix:
        return prefix.group(1)
    match = re.search(r"(\d+)(b?)", stem, re.IGNORECASE)
    if match:
        return f"{match.group(1)}{match.group(2).lower()}"
    return "0"


# ── Title extraction (DOCX, unchanged from v1) ────────────────────────────────

def extract_title(doc: Document, fallback_name: str) -> str:
    """Look for a Heading-style paragraph; fall back to first line or filename."""
    for para in doc.paragraphs:
        if para.style and para.style.name.startswith("Heading") and para.text.strip():
            return para.text.strip()
    for para in doc.paragraphs:
        if para.text.strip():
            return para.text.strip()
    return _strip_order_prefix(Path(fallback_name).stem)


# ── Image extraction (DOCX only, unchanged from v1) ───────────────────────────

def extract_images(docx_path: Path, chapter_id: str) -> list[dict]:
    """
    .docx files are ZIP archives.  Images live in word/media/ inside them.
    Extracts and saves every image, returns a list of image records.
    """
    images = []
    chapter_img_dir = IMAGES_DIR / f"chapter_{chapter_id}"

    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            media_files = [f for f in z.namelist() if f.startswith("word/media/")]
            if not media_files:
                return []
            chapter_img_dir.mkdir(parents=True, exist_ok=True)
            for i, media_path in enumerate(media_files):
                original_name = Path(media_path).name
                saved_name    = f"ch{chapter_id}_image_{i+1:02d}_{original_name}"
                save_to       = chapter_img_dir / saved_name
                with z.open(media_path) as img_data:
                    save_to.write_bytes(img_data.read())
                images.append({
                    "index":                 i + 1,
                    "saved_filename":        saved_name,
                    "original_name":         original_name,
                    "relative_path":         str(save_to.relative_to(BASE_DIR)),
                    "illustration_prompt":   None,
                    "generated_image_path":  None,
                })
    except Exception as e:
        warn(f"Could not extract images from {docx_path.name}: {e}")

    return images


# ── Paragraph extraction (DOCX, unchanged from v1) ────────────────────────────

def extract_paragraphs(doc: Document) -> tuple[list[dict], str]:
    """Read every paragraph from a python-docx Document object."""
    paragraphs  = []
    text_chunks = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        paragraphs.append({
            "style": para.style.name if para.style else "Normal",
            "text":  text,
        })
        text_chunks.append(text)
    full_text = "\n\n".join(text_chunks)
    return paragraphs, full_text


# ── Format readers (new formats) ──────────────────────────────────────────────

def _read_pdf(path: Path) -> list[dict]:
    """
    Extract text from a PDF using pdfplumber.
    Returns a flat list of paragraph dicts (one per non-empty line).
    Returns [] and warns if pdfplumber is missing or the PDF is image-only.
    """
    pdfplumber = _import_pdfplumber()
    if pdfplumber is None:
        warn(
            f"pdfplumber is not installed — cannot read {path.name}.\n"
            "     Install it with:  pip install pdfplumber"
        )
        return []

    paragraphs = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if not text:
                    continue
                for line in text.splitlines():
                    line = line.strip()
                    if line:
                        paragraphs.append({"style": "Normal", "text": line})
    except Exception as e:
        warn(f"PDF read failed for {path.name}: {e}")
        return []

    if not paragraphs:
        warn(
            f"{path.name} appears to be a scanned / image-only PDF — "
            "no text could be extracted.  OCR is not yet supported; "
            "skipping this file."
        )

    return paragraphs


def _read_epub(path: Path) -> list[tuple[str, list[dict]]]:
    """
    Parse an EPUB file and return a list of (title, paragraphs) — one per
    spine item (i.e. one per chapter as the EPUB author structured it).

    Returns [] and warns if ebooklib / beautifulsoup4 are missing.
    """
    ebooklib, epub_mod, BeautifulSoup = _import_epub_libs()
    if ebooklib is None:
        warn(
            f"ebooklib / beautifulsoup4 not installed — cannot read {path.name}.\n"
            "     Install them with:  pip install ebooklib beautifulsoup4"
        )
        return []

    chapters_out = []
    try:
        book = epub_mod.read_epub(str(path), options={"ignore_ncx": True})

        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            raw_html = item.get_content()
            if not raw_html:
                continue

            soup = BeautifulSoup(raw_html, "html.parser")

            # Skip navigation / TOC pages (very short content)
            body_text = soup.get_text(strip=True)
            if len(body_text) < 120:
                continue

            # Extract chapter title from first heading
            title = ""
            for tag in soup.find_all(["h1", "h2", "h3"]):
                t = tag.get_text(strip=True)
                if t:
                    title = t
                    break

            # Build paragraphs
            paragraphs = []
            for elem in soup.find_all(["p", "h1", "h2", "h3", "h4"]):
                text = elem.get_text(strip=True)
                if not text:
                    continue
                style = "Normal"
                if elem.name in ("h1", "h2"):
                    style = "Heading 1"
                elif elem.name in ("h3", "h4"):
                    style = "Heading 2"
                paragraphs.append({"style": style, "text": text})

            if paragraphs:
                chapters_out.append((title, paragraphs))

    except Exception as e:
        warn(f"EPUB read failed for {path.name}: {e}")

    return chapters_out


def _read_txt(path: Path) -> list[dict]:
    """
    Read a plain-text or Markdown file.
    For .md: headings are converted to style="Heading N" and inline
    markdown (bold, italic, links, code) is stripped to plain text.
    For .txt: each non-empty line becomes a paragraph.
    """
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        warn(f"Could not read {path.name}: {e}")
        return []

    is_markdown = path.suffix.lower() == ".md"
    paragraphs  = []

    for line in text.splitlines():
        line_stripped = line.strip()
        if not line_stripped:
            continue

        style = "Normal"

        if is_markdown:
            # Detect ATX headings
            if line_stripped.startswith("### "):
                style         = "Heading 3"
                line_stripped = line_stripped[4:].strip()
            elif line_stripped.startswith("## "):
                style         = "Heading 2"
                line_stripped = line_stripped[3:].strip()
            elif line_stripped.startswith("# "):
                style         = "Heading 1"
                line_stripped = line_stripped[2:].strip()
            else:
                # Strip inline markdown: **bold**, *italic*, `code`, [link](url)
                line_stripped = re.sub(r"\*{1,2}([^*]+)\*{1,2}", r"\1", line_stripped)
                line_stripped = re.sub(r"`([^`]+)`",              r"\1", line_stripped)
                line_stripped = re.sub(r"\[([^\]]+)\]\([^)]+\)",  r"\1", line_stripped)
                line_stripped = re.sub(r"^[-*>]\s+",              "",    line_stripped)

        if line_stripped:
            paragraphs.append({"style": style, "text": line_stripped})

    return paragraphs


# ── Chapter record builder ─────────────────────────────────────────────────────

def _build_chapter_record(
    chapter_id:  str,
    chapter_num: int | None,
    title:       str,
    source_file: str,
    paragraphs:  list[dict],
    images:      list[dict],
    source_fmt:  str = "docx",
) -> dict:
    """
    Assemble the canonical chapter JSON dict.
    Schema is identical to v1 — downstream agents depend on every field here.
    Only additive change: metadata.source_format (new field, safe to add).
    """
    full_text      = "\n\n".join(p["text"] for p in paragraphs)
    word_count     = len(full_text.split())
    sentence_count = len(re.findall(r"[.!?]+", full_text))

    return {
        # ── Identity ──────────────────────────────────────────────
        "chapter_number":   chapter_num,
        "chapter_id":       chapter_id,
        "title":            title,
        "source_file":      source_file,

        # ── Stats ─────────────────────────────────────────────────
        "word_count":       word_count,
        "sentence_count":   sentence_count,
        "paragraph_count":  len(paragraphs),

        # ── Content ───────────────────────────────────────────────
        "paragraphs":       paragraphs,
        "full_text":        full_text,

        # ── Images ────────────────────────────────────────────────
        "images":           images,

        # ── Pipeline status (unchanged schema) ─────────────────────
        "pipeline_status": {
            "ingested":                      True,
            "consistency_checked":           False,
            "consistency_issues":            [],
            "editing_complete":              False,
            "editing_creativity_level":      None,
            "illustration_prompt_generated": False,
            "formatted":                     False,
            "qc_passed":                     False,
            "qc_issues":                     [],
        },

        # ── Metadata ──────────────────────────────────────────────
        "metadata": {
            "ingested_at":      datetime.now().isoformat(),
            "pipeline_version": "1.1",
            "source_format":    source_fmt,       # additive — safe for v1 consumers
        },
    }


# ── Single-file chapter splitter ──────────────────────────────────────────────

def _heading_number(text: str) -> int | None:
    """Pull an integer out of a heading like 'Chapter 3' or '12.'"""
    m = re.search(r"\d+", text)
    return int(m.group()) if m else None


def _find_break_indices(paragraphs: list[dict]) -> list[int]:
    """
    Return paragraph indices that look like chapter-start headings.
    We trigger on Heading 1 style OR a text pattern match.
    """
    return [
        i for i, p in enumerate(paragraphs)
        if p["style"].startswith("Heading 1") or _is_chapter_marker(p["text"])
    ]


def _split_paragraphs(
    paragraphs:    list[dict],
    break_indices: list[int],
    source_name:   str,
) -> list[tuple[str, list[dict], int | None]]:
    """
    Given a flat paragraph list and the indices of chapter-start headings,
    slice it into chapters.

    Returns a list of (title, paragraphs_for_that_chapter, chapter_number_or_None).
    The heading paragraph is kept INSIDE paragraphs (consistent with v1 behaviour
    where the title heading is both chapter_data.title AND in chapter_data.paragraphs).
    """
    results = []

    # Content before the first chapter marker → treat as front matter if substantial
    if break_indices[0] > 0:
        pre_paras = paragraphs[: break_indices[0]]
        pre_words = sum(len(p["text"].split()) for p in pre_paras)
        if pre_words > 50:
            results.append(("Front Matter", pre_paras, None))

    for idx, start in enumerate(break_indices):
        end          = break_indices[idx + 1] if idx + 1 < len(break_indices) else len(paragraphs)
        chapter_paras = paragraphs[start:end]

        if not chapter_paras:
            continue

        title  = chapter_paras[0]["text"]
        ch_num = _heading_number(title)
        results.append((title, chapter_paras, ch_num))

    return results


def _build_split_records(
    splits:      list[tuple[str, list[dict], int | None]],
    source_name: str,
    source_fmt:  str,
) -> list[dict]:
    """Turn the list of split chapters into full chapter record dicts."""
    records = []
    for seq, (title, paras, ch_num) in enumerate(splits, start=1):
        # Prefer the number detected from the heading; fall back to sequence position
        num   = ch_num if ch_num is not None else seq
        ch_id = "epilogue" if "epilogue" in title.lower() else f"{num:02d}"
        records.append(
            _build_chapter_record(
                chapter_id  = ch_id,
                chapter_num = None if ch_id == "epilogue" else num,
                title       = title,
                source_file = source_name,
                paragraphs  = paras,
                images      = [],        # images not attributable after split
                source_fmt  = source_fmt,
            )
        )
    return records


# ── Per-file processor ────────────────────────────────────────────────────────

def process_file(path: Path) -> list[dict]:
    """
    Read one input file and return a list of chapter record dicts.

    For per-chapter files (the normal case) this list has exactly one item.
    For single-file manuscripts that contain multiple chapters, the list will
    have one item per detected chapter.

    Returns [] on failure (caller handles the error).
    """
    ext        = path.suffix.lower()
    source_name = path.name

    # ── DOCX / DOC ────────────────────────────────────────────────────────────
    if ext in (".docx", ".doc"):
        doc = Document(str(path))
        paragraphs, _full_text = extract_paragraphs(doc)

        if not paragraphs:
            warn(f"No text found in {source_name}")
            return []

        breaks = _find_break_indices(paragraphs)

        if len(breaks) >= 2:
            # Multiple chapter headings → single-file manuscript, split it
            info(f"  → {len(breaks)} chapter markers found — splitting into chapters")
            splits = _split_paragraphs(paragraphs, breaks, source_name)
            return _build_split_records(splits, source_name, "docx")
        else:
            # Standard per-chapter file — use existing v1 logic exactly
            chapter_id  = extract_chapter_id(source_name)
            chapter_num = extract_chapter_number(source_name)
            title       = extract_title(doc, source_name)
            images      = extract_images(path, chapter_id)
            return [
                _build_chapter_record(
                    chapter_id  = chapter_id,
                    chapter_num = chapter_num,
                    title       = title,
                    source_file = source_name,
                    paragraphs  = paragraphs,
                    images      = images,
                    source_fmt  = "docx",
                )
            ]

    # ── PDF ───────────────────────────────────────────────────────────────────
    elif ext == ".pdf":
        paragraphs = _read_pdf(path)
        if not paragraphs:
            return []

        breaks = _find_break_indices(paragraphs)

        if len(breaks) >= 2:
            info(f"  → {len(breaks)} chapter markers found — splitting into chapters")
            splits = _split_paragraphs(paragraphs, breaks, source_name)
            return _build_split_records(splits, source_name, "pdf")
        else:
            chapter_id  = extract_chapter_id(source_name)
            chapter_num = extract_chapter_number(source_name)
            # Title: first heading paragraph, or first paragraph, or filename
            title = next(
                (p["text"] for p in paragraphs if p["style"].startswith("Heading")),
                paragraphs[0]["text"] if paragraphs else _strip_order_prefix(Path(source_name).stem),
            )
            return [
                _build_chapter_record(
                    chapter_id  = chapter_id,
                    chapter_num = chapter_num,
                    title       = title,
                    source_file = source_name,
                    paragraphs  = paragraphs,
                    images      = [],
                    source_fmt  = "pdf",
                )
            ]

    # ── EPUB ──────────────────────────────────────────────────────────────────
    elif ext == ".epub":
        epub_chapters = _read_epub(path)
        if not epub_chapters:
            return []

        # EPUB spine already gives us per-chapter content
        base_num = extract_chapter_number(source_name) or 1
        records  = []
        for i, (title, paras) in enumerate(epub_chapters):
            ch_num = base_num + i
            ch_id  = "epilogue" if "epilogue" in title.lower() else f"{ch_num:02d}"
            records.append(
                _build_chapter_record(
                    chapter_id  = ch_id,
                    chapter_num = None if ch_id == "epilogue" else ch_num,
                    title       = title or f"Chapter {ch_num}",
                    source_file = source_name,
                    paragraphs  = paras,
                    images      = [],
                    source_fmt  = "epub",
                )
            )
        return records

    # ── TXT / MARKDOWN ────────────────────────────────────────────────────────
    elif ext in (".txt", ".md"):
        paragraphs = _read_txt(path)
        if not paragraphs:
            warn(f"No content found in {source_name}")
            return []

        breaks = _find_break_indices(paragraphs)

        if len(breaks) >= 2:
            info(f"  → {len(breaks)} chapter markers found — splitting into chapters")
            splits = _split_paragraphs(paragraphs, breaks, source_name)
            return _build_split_records(splits, source_name, ext.lstrip("."))
        else:
            chapter_id  = extract_chapter_id(source_name)
            chapter_num = extract_chapter_number(source_name)
            title = next(
                (p["text"] for p in paragraphs if p["style"].startswith("Heading")),
                _strip_order_prefix(Path(source_name).stem),
            )
            return [
                _build_chapter_record(
                    chapter_id  = chapter_id,
                    chapter_num = chapter_num,
                    title       = title,
                    source_file = source_name,
                    paragraphs  = paragraphs,
                    images      = [],
                    source_fmt  = ext.lstrip("."),
                )
            ]

    else:
        warn(f"Unsupported file type '{ext}' — skipping {source_name}")
        return []


# ── Setup ─────────────────────────────────────────────────────────────────────

def setup_directories():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    IMAGES_DIR.mkdir(parents=True, exist_ok=True)
    ok("Output directories ready.")


def save_json(data: dict, path: Path):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


def _get_book_title() -> str:
    """Read the book title from config.json if available, otherwise generic."""
    config_path = BASE_DIR / "config.json"
    if config_path.exists():
        try:
            with open(config_path) as f:
                cfg = json.load(f)
            return cfg.get("book", {}).get("title", "")
        except Exception:
            pass
    return ""


# ── Main ──────────────────────────────────────────────────────────────────────

def run():
    print()
    print("═" * 62)
    print("  AGENT 1 — INGESTION")
    book_title = _get_book_title()
    if book_title:
        print(f"  {book_title}")
    print("═" * 62)
    print()

    setup_directories()
    print()

    # Accepted input formats
    ACCEPTED = {".docx", ".doc", ".pdf", ".epub", ".txt", ".md"}
    all_files = sorted(
        f for f in INPUT_DIR.glob("*")
        if f.is_file() and f.suffix.lower() in ACCEPTED
    )

    if not all_files:
        warn(f"No supported files found in:  {INPUT_DIR}")
        warn("Supported formats: .docx, .pdf, .epub, .txt, .md")
        return

    info(f"Found {len(all_files)} input file(s). Starting...\n")

    results  = []
    failures = []

    for file_path in all_files:
        info(f"Reading: {file_path.name}")
        try:
            chapters = process_file(file_path)

            if not chapters:
                warn(f"No content extracted from {file_path.name} — skipping.\n")
                failures.append({"file": file_path.name, "error": "No content extracted"})
                continue

            if len(chapters) > 1:
                info(f"  → Produced {len(chapters)} chapter(s) from this file")

            for chapter_data in chapters:
                ch_id = chapter_data["chapter_id"]

                if ch_id == "epilogue":
                    output_path = OUTPUT_DIR / "epilogue.json"
                else:
                    padded      = ch_id.zfill(2) if ch_id.isdigit() else ch_id
                    output_path = OUTPUT_DIR / f"chapter_{padded}.json"

                # Collision protection: if this chapter file already exists
                # from a previous file in this batch, rename with a suffix
                if output_path.exists():
                    existing = json.loads(output_path.read_text(encoding="utf-8"))
                    existing_source = existing.get("source_file", "unknown")
                    warn(f"  ⚠ chapter_{padded}.json already exists (from {existing_source})")
                    suffix = "b"
                    while True:
                        alt_path = OUTPUT_DIR / f"chapter_{padded}_{suffix}.json"
                        if not alt_path.exists():
                            break
                        suffix = chr(ord(suffix) + 1)  # b → c → d → ...
                    output_path = alt_path
                    chapter_data["chapter_id"] = f"{padded}_{suffix}"
                    warn(f"    → Saving as {output_path.name} instead (rename in Chapter Manager if needed)")

                save_json(chapter_data, output_path)
                ok(f"Saved → {output_path.name}")

                print(
                    f"     Chapter {ch_id}: \"{chapter_data['title']}\"\n"
                    f"     {chapter_data['word_count']:,} words  |  "
                    f"{chapter_data['paragraph_count']} paragraphs  |  "
                    f"{len(chapter_data['images'])} image(s)"
                )
                print()

                results.append({
                    "chapter_number": chapter_data["chapter_number"],
                    "chapter_id":     ch_id,
                    "title":          chapter_data["title"],
                    "word_count":     chapter_data["word_count"],
                    "image_count":    len(chapter_data["images"]),
                    "json_file":      output_path.name,
                    "status":         "success",
                })

        except Exception as e:
            err(f"Failed on {file_path.name}: {e}\n")
            failures.append({"file": file_path.name, "error": str(e)})

    # ── Write summary file ─────────────────────────────────────────────────────
    summary = {
        "total_chapters": len(results),
        "total_words":    sum(r["word_count"]  for r in results),
        "total_images":   sum(r["image_count"] for r in results),
        "failures":       failures,
        "chapters":       results,
        "completed_at":   datetime.now().isoformat(),
    }
    if book_title:
        summary["book_title"] = book_title

    save_json(summary, OUTPUT_DIR / "ingestion_summary.json")

    # ── Final report ───────────────────────────────────────────────────────────
    print("═" * 62)
    print()
    if results:
        ok("INGESTION COMPLETE")
        print(f"     Chapters processed : {len(results)}")
        print(f"     Total word count   : {summary['total_words']:,}")
        print(f"     Embedded images    : {summary['total_images']}")
        print(f"     Summary file       : output/ingested/ingestion_summary.json")
    if failures:
        print()
        warn(f"{len(failures)} file(s) could not be processed:")
        for f in failures:
            err(f"  {f['file']}: {f['error']}")
    print()
    print("  Next step: run Agent 2 (Consistency Check)")
    print("═" * 62)
    print()


if __name__ == "__main__":
    run()
